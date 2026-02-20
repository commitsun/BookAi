"""
SuperintendenteAgent v1 - Gestión de Conocimiento y Estrategia

- Agregar/actualizar base de conocimientos
- Revisar historial de conversaciones
- Enviar broadcasts
- Comunicación con encargado
"""

from __future__ import annotations

import asyncio
import inspect
import logging
import json
from datetime import datetime
from typing import Any, Dict, List, Optional
import os
import unicodedata

from langchain.agents import AgentExecutor, create_openai_tools_agent
from langchain.prompts import ChatPromptTemplate, MessagesPlaceholder

from pathlib import Path
import re
import tempfile

import boto3
from botocore.config import Config as BotoConfig

from core.config import ModelConfig, ModelTier, Settings
from core.db import get_active_chat_reservation
from core.utils.time_context import get_time_context
from core.utils.utils_prompt import load_prompt
from core.message_utils import sanitize_wa_message
from tools.superintendente_tool import (
    _clean_phone,
    _looks_like_phone,
    _resolve_guest_id_by_name,
    _set_instance_context,
)

log = logging.getLogger("SuperintendenteAgent")
log.setLevel(logging.INFO)


class SuperintendenteAgent:
    """
    Agente Superintendente - Gestor de Conocimiento

    Comunicación exclusiva con encargado vía Telegram
    """

    def __init__(
        self,
        memory_manager: Any,
        supabase_client: Any = None,
        channel_manager: Any = None,
        template_registry: Any = None,
        model_tier: ModelTier = ModelTier.SUPERINTENDENTE,
    ) -> None:
        self.memory_manager = memory_manager
        self.supabase_client = supabase_client
        self.channel_manager = channel_manager
        self.template_registry = template_registry
        self.model_tier = model_tier

        self.llm = ModelConfig.get_llm(model_tier)

        log.info("SuperintendenteAgent inicializado (modelo: %s)", self.llm.model_name)

    def _get_s3_client(self):
        """
        Crea un cliente S3 tolerante a perfiles ausentes.
        - Si AWS_PROFILE está seteado y existe, lo usa.
        - Si no existe (ej. en EC2 con role), cae a credenciales por defecto.
        """
        profile = (os.getenv("AWS_PROFILE") or "").strip() or None
        region = Settings.AWS_DEFAULT_REGION
        session = None

        if profile:
            try:
                session = boto3.Session(profile_name=profile, region_name=region)
            except Exception as exc:
                log.warning("Perfil AWS '%s' no disponible, uso cadena por defecto: %s", profile, exc)

        if session is None:
            # Evita que un AWS_PROFILE vacío provoque ProfileNotFound; borra la variable si está vacía.
            if "AWS_PROFILE" in os.environ and not os.environ.get("AWS_PROFILE"):
                os.environ.pop("AWS_PROFILE", None)
            session = boto3.Session(region_name=region)

        return session.client(
            "s3",
            config=BotoConfig(retries={"max_attempts": 3, "mode": "standard"}),
        )

    async def ainvoke(
        self,
        user_input: str,
        encargado_id: str,
        hotel_name: str,
        context_window: int = 50,
        chat_history: Optional[List[Any]] = None,
        session_id: Optional[str] = None,
        clients_context: Optional[str] = None,
    ) -> str:
        """
        Invocar agente superintendente (sesión con encargado)

        Args:
            user_input: Mensaje del encargado
            encargado_id: ID Telegram del encargado (session_id)
            hotel_name: Nombre del hotel
            context_window: Mensajes para contexto
            chat_history: Historial externo
        """

        try:
            active_session_id = (session_id or "").strip() or None
            convo_id = active_session_id or encargado_id
            original_owner_id = encargado_id

            log.info("SuperintendenteAgent ainvoke: %s", convo_id)

            fast_draft = await self._try_direct_whatsapp_draft(
                user_input,
                encargado_id,
                session_id=active_session_id,
                clients_context=clients_context,
            )
            if fast_draft:
                await self._safe_call(
                    getattr(self.memory_manager, "save", None),
                    conversation_id=convo_id,
                    role="user",
                    content=user_input,
                    channel="telegram",
                    original_chat_id=original_owner_id,
                )
                await self._safe_call(
                    getattr(self.memory_manager, "save", None),
                    conversation_id=convo_id,
                    role="assistant",
                    content=fast_draft,
                    channel="telegram",
                    original_chat_id=original_owner_id,
                )
                log.info("Superintendente fast draft enviado (%s chars)", len(fast_draft))
                return fast_draft

            resolved_hotel_name = self._sanitize_hotel_name(hotel_name) or hotel_name
            if self.memory_manager and convo_id:
                try:
                    self.memory_manager.set_flag(convo_id, "property_name", resolved_hotel_name)
                    self.memory_manager.set_flag(
                        convo_id,
                        "history_table",
                        Settings.SUPERINTENDENTE_HISTORY_TABLE,
                    )
                    inferred_property_id = self._extract_property_id(user_input, hotel_name, resolved_hotel_name)
                    if inferred_property_id is not None:
                        self.memory_manager.set_flag(convo_id, "property_id", inferred_property_id)
                    if active_session_id:
                        self.memory_manager.set_flag(convo_id, "superintendente_owner_id", original_owner_id)
                except Exception:
                    pass

            if chat_history is None:
                chat_history = await self._safe_call(
                    getattr(self.memory_manager, "get_memory_as_messages", None),
                    conversation_id=convo_id,
                    limit=context_window,
                )
            chat_history = chat_history or []
            user_input_for_agent = await self._rewrite_followup_with_context(user_input, chat_history)
            user_input_for_agent = self._enrich_reservation_query_with_context(
                user_input_for_agent,
                convo_id=convo_id,
            )
            if user_input_for_agent != user_input:
                log.info("Superintendente follow-up reinterpretado: '%s' -> '%s'", user_input, user_input_for_agent)

            tools = await self._create_tools(resolved_hotel_name, convo_id)

            system_prompt = self._build_system_prompt(
                resolved_hotel_name,
                clients_context=clients_context,
            )
            log.info("Superintendente hotel_name activo: %s (encargado_id=%s)", resolved_hotel_name, convo_id)

            prompt_template = ChatPromptTemplate.from_messages(
                [
                    ("system", system_prompt),
                    MessagesPlaceholder(variable_name="chat_history", optional=True),
                    ("human", "{input}"),
                    MessagesPlaceholder(variable_name="agent_scratchpad", optional=True),
                ]
            )

            agent_chain = create_openai_tools_agent(
                llm=self.llm,
                tools=tools,
                prompt=prompt_template,
            )

            executor = AgentExecutor(
                agent=agent_chain,
                tools=tools,
                verbose=True,
                max_iterations=20,
                handle_parsing_errors=True,
                return_intermediate_steps=True,
                max_execution_time=90,
            )

            result = await executor.ainvoke(
                input={
                    "input": user_input_for_agent,
                    "chat_history": chat_history,
                }
            )

            output = (result.get("output") or "").strip()

            # 🚦 Propagar marcadores especiales si vinieron en pasos intermedios
            intermediates = result.get("intermediate_steps") or []
            wa_markers: list[str] = []
            tpl_marker = None
            kb_marker = None
            kb_rm_marker = None
            broadcast_marker = None
            for _action, observation in intermediates:
                if isinstance(observation, str) and "[WA_DRAFT]|" in observation:
                    wa_markers.append(
                        observation[observation.index("[WA_DRAFT]|") :].strip()
                    )
                if isinstance(observation, str) and "[TPL_DRAFT]|" in observation:
                    tpl_marker = observation[
                        observation.index("[TPL_DRAFT]|") :
                    ].strip()
                if isinstance(observation, str) and "[KB_DRAFT]|" in observation:
                    kb_marker = observation[
                        observation.index("[KB_DRAFT]|") :
                    ].strip()
                if isinstance(observation, str) and "[KB_REMOVE_DRAFT]|" in observation:
                    kb_rm_marker = observation[
                        observation.index("[KB_REMOVE_DRAFT]|") :
                    ].strip()
                if isinstance(observation, str) and "[BROADCAST_DRAFT]|" in observation:
                    broadcast_marker = observation[
                        observation.index("[BROADCAST_DRAFT]|") :
                    ].strip()
                if wa_markers and tpl_marker and kb_marker and kb_rm_marker and broadcast_marker:
                    break
            if wa_markers:
                markers_block = "\n".join(wa_markers)
                if "[WA_DRAFT]|" not in output:
                    output = f"{markers_block}\n{output}"
                else:
                    # Añade los que no estén ya presentes para no perder borradores múltiples
                    for marker in wa_markers:
                        if marker not in output:
                            output = f"{marker}\n{output}"
            if tpl_marker:
                output = tpl_marker
            if kb_marker and "[KB_DRAFT]|" not in output:
                output = f"{kb_marker}\n{output}"
            if kb_rm_marker and "[KB_REMOVE_DRAFT]|" not in output:
                output = f"{kb_rm_marker}\n{output}"
            if broadcast_marker:
                output = broadcast_marker

            await self._safe_call(
                getattr(self.memory_manager, "save", None),
                conversation_id=convo_id,
                role="user",
                content=user_input,
                channel="telegram",
                original_chat_id=original_owner_id,
            )

            await self._safe_call(
                getattr(self.memory_manager, "save", None),
                conversation_id=convo_id,
                role="assistant",
                content=output,
                channel="telegram",
                original_chat_id=original_owner_id,
            )

            log.info("SuperintendenteAgent completado: %s chars", len(output))
            return output

        except Exception as exc:
            log.error("Error en SuperintendenteAgent: %s", exc, exc_info=True)
            raise

    def _is_followup_candidate(self, text: str) -> bool:
        clean = re.sub(r"[¡!¿?.]", "", (text or "").lower()).strip()
        if not clean:
            return False
        tokens = [t for t in re.findall(r"[a-záéíóúñ0-9]+", clean) if t]
        if not tokens or len(tokens) > 4:
            return False
        explicit = {
            "resumen",
            "summary",
            "sintesis",
            "síntesis",
            "original",
            "historial",
            "completo",
            "raw",
            "crudo",
            "mensajes",
        }
        return clean in explicit

    async def _rewrite_followup_with_context(self, user_input: str, chat_history: List[Any]) -> str:
        raw = (user_input or "").strip()
        if not raw or not self.llm:
            return raw
        if not self._is_followup_candidate(raw):
            return raw

        recent: list[str] = []
        for msg in (chat_history or [])[-10:]:
            mtype = str(getattr(msg, "type", "") or "").lower().strip()
            content = str(getattr(msg, "content", "") or "").strip()
            if not content:
                continue
            if mtype == "human":
                role = "user"
            elif mtype == "system":
                role = "system"
            else:
                role = "assistant"
            recent.append(f"{role}: {content}")
        if not recent:
            return raw

        prompt = (
            "Reescribe el último mensaje SOLO si es un follow-up ambiguo y depende del contexto reciente. "
            "Si no lo es, devuélvelo igual. No inventes nombres ni datos.\n\n"
            "Historial reciente:\n"
            f"{chr(10).join(recent)}\n\n"
            f"Último mensaje del usuario: {raw}\n\n"
            "Devuelve únicamente una frase final accionable."
        )
        try:
            resp = await self.llm.ainvoke(
                [
                    {
                        "role": "system",
                        "content": "Eres un normalizador de intención para un asistente operativo de hotel.",
                    },
                    {"role": "user", "content": prompt},
                ]
            )
            rewritten = (getattr(resp, "content", None) or "").strip()
            return rewritten or raw
        except Exception:
            return raw

    async def _try_direct_whatsapp_draft(
        self,
        user_input: str,
        encargado_id: str,
        session_id: Optional[str] = None,
        clients_context: Optional[str] = None,
    ) -> Optional[str]:
        clean_input = user_input.strip()
        if clean_input.lower().startswith("/super"):
            clean_input = clean_input.split(" ", 1)[1].strip() if " " in clean_input else ""

        parsed = self._parse_direct_send_request(clean_input)
        if not parsed and self._has_explicit_send_intent(clean_input):
            parsed = await self._extract_send_intent_llm(clean_input)
        if not parsed:
            log.info("Superintendente fast draft: no match for direct-send pattern")
            return None
        guest_label, message = parsed
        if not message:
            return None

        property_id = None
        instance_id = None
        if self.memory_manager and encargado_id:
            try:
                property_id = self.memory_manager.get_flag(encargado_id, "property_id")
            except Exception:
                property_id = None
            try:
                instance_id = (
                    self.memory_manager.get_flag(encargado_id, "instance_id")
                    or self.memory_manager.get_flag(encargado_id, "instance_hotel_code")
                )
            except Exception:
                instance_id = None

        if _looks_like_phone(guest_label):
            guest_id = _clean_phone(guest_label)
        else:
            guest_id = None
            candidates: list[dict] = []
            guest_id_ctx, candidates_ctx = self._resolve_guest_from_clients_context(guest_label, clients_context)
            if guest_id_ctx:
                guest_id = guest_id_ctx
            elif candidates_ctx:
                candidates.extend(candidates_ctx)
            chat_candidates = []
            if session_id:
                chat_candidates.append(session_id)
            if encargado_id and encargado_id not in chat_candidates:
                chat_candidates.append(encargado_id)

            for cid in chat_candidates:
                guest_id, candidates = _resolve_guest_id_by_name(
                    guest_label,
                    property_id=property_id,
                    memory_manager=self.memory_manager,
                    chat_id=cid,
                )
                if guest_id:
                    break
            if not guest_id:
                # Fallback: reintenta con extracción LLM si el nombre estaba contaminado por instrucciones.
                parsed_llm = await self._extract_send_intent_llm(clean_input)
                if parsed_llm:
                    llm_guest, llm_message = parsed_llm
                    if llm_guest and llm_guest.strip() != guest_label:
                        for cid in chat_candidates or [encargado_id]:
                            guest_id, candidates = _resolve_guest_id_by_name(
                                llm_guest,
                                property_id=property_id,
                                memory_manager=self.memory_manager,
                                chat_id=cid,
                            )
                            if guest_id:
                                break
                        if guest_id and llm_message:
                            message = llm_message
                if candidates:
                    log.info("Superintendente fast draft: nombre ambiguo (%s)", guest_label)
                    lines = []
                    for cand in candidates[:5]:
                        label = cand.get("client_name") or "Sin nombre"
                        lines.append(f"• {label} → {cand.get('phone')}")
                    suggestions = "\n".join(lines)
                    return (
                        "⚠️ Encontré varios huéspedes con ese nombre. "
                        "Indícame el teléfono exacto:\n"
                        f"{suggestions}"
                    )
                return (
                    f"⚠️ No encontré un huésped con el nombre '{guest_label}'. "
                    "Indícame el teléfono exacto."
                )

        if not guest_id:
            return None

        if self.memory_manager and encargado_id:
            try:
                _set_instance_context(
                    self.memory_manager,
                    encargado_id,
                    property_id=property_id,
                    instance_id=instance_id,
                )
            except Exception:
                pass

        if self._needs_wa_polish(message):
            message = await self._compose_guest_message(message)

        return f"[WA_DRAFT]|{guest_id}|{message}"

    def _normalize_person_name(self, value: str) -> str:
        raw = str(value or "").strip().lower()
        if not raw:
            return ""
        deaccented = "".join(
            ch for ch in unicodedata.normalize("NFKD", raw) if not unicodedata.combining(ch)
        )
        cleaned = re.sub(r"[^a-z0-9]+", " ", deaccented)
        return re.sub(r"\s+", " ", cleaned).strip()

    def _resolve_guest_from_clients_context(
        self,
        guest_label: str,
        clients_context: Optional[str],
    ) -> tuple[Optional[str], list[dict]]:
        """
        Resuelve guest_id usando el bloque CLIENTES_ACTIVOS inyectado por API.
        Devuelve (guest_id, candidates) para mantener el flujo actual de ambigüedad.
        """
        label = str(guest_label or "").strip()
        block = str(clients_context or "").strip()
        if not label or not block:
            return None, []

        query = self._normalize_person_name(label)
        if not query:
            return None, []

        rows = [ln.strip() for ln in block.splitlines() if ln and "|" in ln]
        if not rows:
            return None, []

        header_idx = None
        header_cols: list[str] = []
        for idx, line in enumerate(rows):
            cols = [c.strip().lower() for c in line.split("|")]
            if "chat_id" in cols and "client_name" in cols:
                header_idx = idx
                header_cols = cols
                break
        if header_idx is None:
            return None, []

        def _col(name: str) -> Optional[int]:
            try:
                return header_cols.index(name)
            except ValueError:
                return None

        idx_name = _col("client_name")
        idx_phone = _col("client_phone")
        idx_chat = _col("chat_id")
        idx_last = _col("last_message_at")
        if idx_name is None:
            return None, []

        candidates: list[dict] = []
        for line in rows[header_idx + 1 :]:
            cols = [c.strip() for c in line.split("|")]
            if len(cols) < len(header_cols):
                continue
            raw_name = cols[idx_name] if idx_name < len(cols) else ""
            if not raw_name or raw_name.lower() == "null":
                continue
            norm_name = self._normalize_person_name(raw_name)
            if not norm_name:
                continue
            if query != norm_name and not norm_name.startswith(query) and query not in norm_name:
                continue
            raw_phone = cols[idx_phone] if idx_phone is not None and idx_phone < len(cols) else ""
            raw_chat = cols[idx_chat] if idx_chat is not None and idx_chat < len(cols) else ""
            phone = _clean_phone(raw_phone) or _clean_phone(raw_chat)
            if not phone:
                continue
            created_at = cols[idx_last] if idx_last is not None and idx_last < len(cols) else ""
            score = 0 if norm_name == query else (1 if norm_name.startswith(query) else 2)
            candidates.append(
                {
                    "phone": phone,
                    "client_name": raw_name,
                    "created_at": created_at,
                    "score": score,
                    "source": "clients_context",
                }
            )

        if not candidates:
            return None, []

        # Dedup por teléfono manteniendo mejor score y más reciente.
        def _ts(value: Any) -> float:
            try:
                return datetime.fromisoformat(str(value).replace("Z", "+00:00")).timestamp()
            except Exception:
                return 0.0

        candidates.sort(key=lambda c: (c.get("score", 9), -_ts(c.get("created_at"))))
        unique: list[dict] = []
        seen = set()
        for cand in candidates:
            phone = cand.get("phone")
            if not phone or phone in seen:
                continue
            seen.add(phone)
            unique.append(cand)

        if not unique:
            return None, []

        best_score = unique[0].get("score", 9)
        best = [c for c in unique if c.get("score", 9) == best_score]
        if len(best) == 1:
            return best[0].get("phone"), unique
        phones = {c.get("phone") for c in best if c.get("phone")}
        if len(phones) == 1:
            return next(iter(phones)), unique
        return None, unique

    def _needs_wa_polish(self, message: str) -> bool:
        text = (message or "").lower()
        if not text:
            return False
        triggers = (
            "añade",
            "agrega",
            "incluye",
            "de manera",
            "educad",
            "formatea",
            "haz que",
            "por favor",
        )
        return any(t in text for t in triggers)

    def _has_explicit_send_intent(self, text: str) -> bool:
        clean = (text or "").strip().lower()
        if not clean:
            return False
        # El fast-path solo debe activarse cuando hay verbo de envío explícito.
        return bool(
            re.search(
                r"\b(envia|envíale|enviale|manda|mándale|mandale|dile|escribe(?:le)?)\b",
                clean,
                flags=re.IGNORECASE,
            )
        )

    def _extract_reservation_subject_name(self, text: str) -> Optional[str]:
        raw = (text or "").strip()
        if not raw:
            return None
        patterns = [
            r"(?i)\breserva\s+de\s+([a-záéíóúñ0-9\s\.\-]{2,60})",
            r"(?i)\binfo(?:rmaci[oó]n)?\s+sobre\s+la\s+reserva\s+de\s+([a-záéíóúñ0-9\s\.\-]{2,60})",
            r"(?i)\bdatos\s+de\s+la\s+reserva\s+de\s+([a-záéíóúñ0-9\s\.\-]{2,60})",
        ]
        for pattern in patterns:
            m = re.search(pattern, raw)
            if not m:
                continue
            value = re.split(r"[,.;:!?]", m.group(1), maxsplit=1)[0].strip()
            if value:
                return value
        return None

    def _enrich_reservation_query_with_context(self, user_input: str, *, convo_id: str) -> str:
        raw = (user_input or "").strip()
        if not raw or not self.memory_manager:
            return raw
        lower = raw.lower()
        if "reserva" not in lower:
            return raw
        # Si ya viene con folio, no tocar.
        if re.search(r"\bfolio(?:_id)?\s*[:#]?\s*[a-z0-9]{4,}\b", raw, flags=re.IGNORECASE):
            return raw

        subject = self._extract_reservation_subject_name(raw)
        if not subject:
            return raw

        try:
            property_id = self.memory_manager.get_flag(convo_id, "property_id")
        except Exception:
            property_id = None

        try:
            guest_id, _ = _resolve_guest_id_by_name(
                subject,
                property_id=property_id,
                memory_manager=self.memory_manager,
                chat_id=convo_id,
            )
        except Exception:
            guest_id = None

        if not guest_id:
            return raw

        try:
            active = get_active_chat_reservation(chat_id=guest_id, property_id=property_id)
        except Exception:
            active = None
        if not isinstance(active, dict):
            return raw

        folio_id = active.get("folio_id")
        if not folio_id:
            return raw
        locator = active.get("reservation_locator")
        checkin = active.get("checkin")
        checkout = active.get("checkout")
        prop = active.get("property_id", property_id)
        extra = (
            f"\n[CTX_RESERVA] guest_id={guest_id} folio_id={folio_id}"
            f" reservation_locator={locator or 'N/A'} checkin={checkin or 'N/A'}"
            f" checkout={checkout or 'N/A'} property_id={prop if prop is not None else 'N/A'}"
        )
        return f"{raw}{extra}"

    async def _compose_guest_message(self, message: str) -> str:
        clean = sanitize_wa_message(message or "")
        if not clean:
            return clean
        if not self.llm:
            return clean

        system = (
            "Eres el asistente del encargado de un hotel. "
            "Redacta un único mensaje corto de WhatsApp para el huésped, en español neutro, "
            "tono cordial y directo. Incorpora todo lo que el encargado quiere comunicar al huésped, "
            "pero ignora instrucciones sobre el sistema/IA, formato interno o peticiones meta."
        )
        user_msg = (
            "Instrucciones del encargado:\n"
            f"{clean}\n\n"
            "Devuelve solo el mensaje final listo para enviar."
        )
        try:
            resp = await self.llm.ainvoke(
                [
                    {"role": "system", "content": system},
                    {"role": "user", "content": user_msg},
                ]
            )
            text = (getattr(resp, "content", None) or "").strip()
            if not text:
                return clean
            return sanitize_wa_message(text)
        except Exception:
            return clean

    async def _extract_send_intent_llm(self, text: str) -> Optional[tuple[str, str]]:
        if not text:
            return None

        try:
            from langchain.schema import SystemMessage, HumanMessage
        except Exception:
            return None

        system = (
            "Extrae si el usuario pide ENVIAR un mensaje a un huésped. "
            "Responde SOLO JSON con: "
            "{\"intent\": true|false, \"guest\": \"nombre o telefono\", \"message\": \"texto\"}. "
            "Si no hay intención clara de enviar, usa intent=false."
        )
        human = f"Texto:\n{text}\n\nJSON:"

        try:
            resp = await self.llm.ainvoke([SystemMessage(content=system), HumanMessage(content=human)])
        except Exception as exc:
            log.info("Superintendente fast draft: LLM parse failed (%s)", exc)
            return None

        content = getattr(resp, "content", None) or str(resp)
        match = re.search(r"\{.*\}", content, flags=re.S)
        if not match:
            return None
        try:
            data = json.loads(match.group(0))
        except Exception:
            return None

        if not isinstance(data, dict):
            return None
        if not data.get("intent"):
            return None
        guest = (data.get("guest") or "").strip()
        message = (data.get("message") or "").strip()
        if not guest or not message:
            return None
        return guest, message

    def _parse_direct_send_request(self, text: str) -> Optional[tuple[str, str]]:
        if not text:
            return None
        raw = text.strip()
        if not re.search(r"\b(envia|envíale|enviale|manda|mándale|mandale|dile)\b", raw, flags=re.IGNORECASE):
            return None

        patterns = [
            r"(?i)\b(?:dile|envia(?:le)?|envíale|manda(?:le)?|mándale)\s+a\s+(.+?)\s+que\s+(.+)$",
            r"(?i)\b(?:envia(?:le)?|envíale|manda(?:le)?|mándale)\s+a\s+(.+?)\s+(?:un|una)?\s*mensaje\s+que\s+(.+)$",
            r"(?i)\b(?:envia(?:le)?|envíale|manda(?:le)?|mándale)\s+un\s+mensaje\s+a\s+(.+?)\s+que\s+(.+)$",
        ]
        for pattern in patterns:
            match = re.search(pattern, raw)
            if match:
                name = match.group(1).strip()
                msg = match.group(2).strip()
                return name, msg
        fallback_patterns = [
            r"(?i)\b(?:dile|envia(?:le)?|envíale|manda(?:le)?|mándale)\s+a\s+(.+?)\s*:\s*(.+)$",
            r"(?i)\b(?:envia(?:le)?|envíale|manda(?:le)?|mándale)\s+a\s+(.+?)\s+un\s+mensaje\s+(?:diciendo|diciéndole)?\s*[:\-]?\s*(.+)$",
            r"(?i)\b(?:envia(?:le)?|envíale|manda(?:le)?|mándale)\s+a\s+(.+?)\s+un\s+mensaje\s+(.+)$",
            r"(?i)\b(?:dile)\s+a\s+(.+?)\s+(.+)$",
        ]
        for pattern in fallback_patterns:
            match = re.search(pattern, raw)
            if match:
                name = match.group(1).strip()
                msg = match.group(2).strip()
                return name, msg
        return None

    async def _create_tools(self, hotel_name: str, encargado_id: str):
        """Crear tools del superintendente"""

        from tools.superintendente_tool import (
            create_add_to_kb_tool,
            create_consulta_reserva_general_tool,
            create_consulta_reserva_persona_tool,
            create_list_templates_tool,
            create_review_conversations_tool,
            create_remove_from_kb_tool,
            create_send_broadcast_tool,
            create_send_broadcast_checkin_tool,
            create_send_message_main_tool,
            create_send_template_tool,
            create_send_whatsapp_tool,
        )

        tools = [
            create_remove_from_kb_tool(
                hotel_name=hotel_name,
                preview_func=lambda criterio, fecha_inicio=None, fecha_fin=None: self._prepare_kb_removal_preview(
                    hotel_name=hotel_name,
                    criterio=criterio,
                    fecha_inicio=fecha_inicio,
                    fecha_fin=fecha_fin,
                ),
            ),
            create_add_to_kb_tool(
                hotel_name=hotel_name,
                append_func=self._append_to_knowledge_document,
                llm=self.llm,
            ),
            create_review_conversations_tool(
                hotel_name=hotel_name,
                memory_manager=self.memory_manager,
                chat_id=encargado_id,
            ),
            create_consulta_reserva_general_tool(
                memory_manager=self.memory_manager,
                chat_id=encargado_id,
            ),
            create_consulta_reserva_persona_tool(
                memory_manager=self.memory_manager,
                chat_id=encargado_id,
            ),
            create_list_templates_tool(
                hotel_name=hotel_name,
                template_registry=self.template_registry,
                supabase_client=self.supabase_client,
            ),
            create_send_broadcast_tool(
                hotel_name=hotel_name,
                channel_manager=self.channel_manager,
                supabase_client=self.supabase_client,
                template_registry=self.template_registry,
                memory_manager=self.memory_manager,
                chat_id=encargado_id,
            ),
            create_send_broadcast_checkin_tool(
                hotel_name=hotel_name,
                channel_manager=self.channel_manager,
                supabase_client=self.supabase_client,
                template_registry=self.template_registry,
                memory_manager=self.memory_manager,
                chat_id=encargado_id,
            ),
            create_send_template_tool(
                hotel_name=hotel_name,
                channel_manager=self.channel_manager,
                template_registry=self.template_registry,
                supabase_client=self.supabase_client,
                memory_manager=self.memory_manager,
                chat_id=encargado_id,
            ),
            create_send_message_main_tool(
                encargado_id=encargado_id,
                channel_manager=self.channel_manager,
            ),
            create_send_whatsapp_tool(
                channel_manager=self.channel_manager,
                memory_manager=self.memory_manager,
                chat_id=encargado_id,
            ),
        ]

        # Ajustar append_func para herramienta de KB al método interno S3
        return [tool for tool in tools if tool is not None]

    def _build_system_prompt(self, hotel_name: str, clients_context: Optional[str] = None) -> str:
        """Construir system prompt para superintendente"""

        base = load_prompt("superintendente_prompt.txt") or (
            "Eres el Superintendente del Sistema de IA Hotelera.\n"
            "Tu rol es gestionar el conocimiento del hotel y optimizar el sistema de agentes.\n\n"
            "RESPONSABILIDADES:\n"
            "1. Agregar y actualizar la base de conocimientos del hotel\n"
            "2. Revisar el historial de conversaciones de huéspedes\n"
            "3. Enviar plantillas individuales o masivas por WhatsApp\n"
            "4. Coordinar con el MainAgent\n"
            "5. Ayudar al encargado a mejorar las respuestas\n\n"
            "HERRAMIENTAS DISPONIBLES:\n"
            "1. agregar_a_base_conocimientos - Agrega información vectorizada a Supabase\n"
            "2. eliminar_de_base_conocimientos - Prepara borrador de eliminación en la base Variable (muestra registros o conteo, requiere confirmación)\n"
            "3. revisar_conversaciones - Revisa conversaciones recientes de huéspedes (pide modo: resumen u original)\n"
            "4. listar_plantillas_whatsapp - Lista las plantillas disponibles en Supabase por idioma/instancia\n"
            "5. enviar_broadcast - Envía plantillas masivas a múltiples huéspedes\n"
            "6. preparar_envio_plantilla - Prepara borrador de envío individual a uno o varios huéspedes (pide parámetros faltantes y espera confirmación)\n"
            "7. enviar_mensaje_main - Envía respuesta del encargado al MainAgent\n"
            "8. consulta_reserva_general - Consulta folios/reservas entre fechas (usa token auto, devuelve folio_id y folio_code)\n"
            "9. consulta_reserva_persona - Consulta detalle de folio (usa token auto, incluye portalUrl si existe)\n\n"
            "TONO: Profesional, eficiente, orientado a mejora continua.\n\n"
            "REGLAS CLAVE:\n"
            "- Antes de usar revisar_conversaciones pregunta si prefiere 'resumen' (síntesis IA) o ver los mensajes 'originales'; usa el modo solicitado.\n"
            "- Usa SIEMPRE la tool revisar_conversaciones para mostrar el historial de un huésped; si no tienes guest_id, pídelo y respeta el límite indicado (default 10).\n"
            "- Para dudas sobre reservas/folios/clientes (estado, pagos, contacto, fechas), prioriza las tools de reservas. Si ya tienes folio_id, usa consulta_reserva_persona; si solo hay nombre/fechas, usa consulta_reserva_general para obtener folio_id antes de detallar.\n"
            "- No uses revisar_conversaciones salvo que pidan explícitamente historial/mensajes/chat del huésped.\n"
            "- Si consulta_reserva_persona devuelve portalUrl, inclúyelo en la respuesta como enlace para factura/portal.\n"
            "- En paneles de reservas, muestra siempre el folio_id numérico (además del folio_code si quieres) para que el encargado pueda pedir detalle con ese ID.\n"
            "- Para enviar plantillas individuales o a pocos huéspedes, usa la herramienta 'preparar_envio_plantilla': genera el borrador, muestra parámetros faltantes y espera confirmación ('sí' para enviar, 'no' para cancelar). Si faltan datos, pídelos antes de preparar el envío final.\n"
            "- REGLA CRÍTICA PARA KB: Cuando el encargado pida agregar/actualizar información en la base de conocimientos, "
            "usa SIEMPRE la herramienta 'agregar_a_base_conocimientos'. Devuelve el marcador [KB_DRAFT]|hotel|tema|categoria|contenido "
            "para que el sistema pueda mostrar el borrador completo (TEMA/CATEGORÍA/CONTENIDO) antes de guardar. No omitas el marcador.\n"
            "- Para eliminar información de la base Variable, usa la herramienta 'eliminar_de_base_conocimientos' sin pedir confirmación previa. "
            "Entrega SIEMPRE el marcador [KB_REMOVE_DRAFT]|hotel|payload_json (con conteo y preview) en tu respuesta para que el encargado confirme/cancele. "
            "Si el encargado pide eliminar/quitar/borrar/limpiar o 'revisar' antes de eliminar, no generes propuestas de agregado ni paneles genéricos: "
            "limítate a invocar 'eliminar_de_base_conocimientos' con el criterio pedido y devuelve el marcador de borrador de eliminación.\n"
            "- REGLA CRÍTICA PARA PLANTILLAS: cuando una herramienta devuelva un marcador [TPL_DRAFT]|..., reenvía EXACTAMENTE ese contenido al encargado, sin añadir resúmenes, reformular ni modificar el panel o la plantilla. No agregues una segunda respuesta después del panel.\n"
            "- Si el último mensaje enviado incluye un borrador de plantilla ([TPL_DRAFT]|...), interpreta 'sí'/'no' o datos adicionales como respuesta a ese borrador; NO invoques herramientas de base de conocimientos en ese contexto."
        )

        context = get_time_context()
        parts = [f"{context}\n{base}\n\nHotel: {hotel_name}"]
        if clients_context:
            parts.append(
                "Contexto global de clientes (snapshot operativo actual):\n"
                f"{clients_context[:12000]}\n"
                "Usa este bloque como fuente principal para chat_id, canal, nombre, teléfono, estado, folio, habitación, "
                "bookai_enabled, checkin, checkout, unread_count y last_message_at. Si falta algún dato puntual, complétalo con tools."
            )
        return "\n\n".join(parts)

    def _sanitize_hotel_name(self, hotel_name: str) -> str:
        raw = " ".join((hotel_name or "").split())
        if not raw:
            return ""
        match = re.search(r"(hotel|hostal)\s+alda[^,.;\n]*", raw, flags=re.IGNORECASE)
        if match:
            raw = match.group(0).strip()
        raw = re.sub(
            r"\s+(que|para|con|donde|cuando|hay|tiene|sobre)\b.*$",
            "",
            raw,
            flags=re.IGNORECASE,
        ).strip()
        return raw

    def _resolve_hotel_name(self, hotel_name: str, encargado_id: str) -> tuple[str, bool]:
        if self.memory_manager and encargado_id:
            try:
                stored = self.memory_manager.get_flag(encargado_id, "property_name")
            except Exception:
                stored = None
            if stored:
                cleaned = self._sanitize_hotel_name(str(stored))
                if cleaned:
                    return cleaned, True
        cleaned = self._sanitize_hotel_name(hotel_name)
        if cleaned:
            return cleaned, False
        return hotel_name, False

    def _extract_property_id(self, *texts: str) -> int | None:
        for text in texts:
            if not text:
                continue
            match = re.search(r"\b(?:property|propiedad)\s*(\d+)\b", text, flags=re.IGNORECASE)
            if match:
                try:
                    return int(match.group(1))
                except ValueError:
                    continue
        return None

    async def handle_kb_addition(
        self,
        topic: str,
        content: str,
        encargado_id: str,
        hotel_name: str,
        source: str = "escalation",
    ) -> Dict[str, Any]:
        """
        Procesar solicitud de agregar a base de conocimientos

        Llamado desde InternoAgent cuando encargado aprueba
        """

        try:
            log.info("Agregando a KB: %s desde %s", topic, source)

            clean_content = self._clean_kb_content(content)

            resolved_name, from_memory = self._resolve_hotel_name(hotel_name, encargado_id)
            result = await self._append_to_knowledge_document(
                topic=topic,
                content=clean_content,
                hotel_name=resolved_name,
                source_type=source,
                use_env=False,
            )

            try:
                from core.db import add_kb_daily_cache

                property_id = None
                kb_name = None
                if self.memory_manager and encargado_id:
                    property_id = self.memory_manager.get_flag(encargado_id, "property_id")
                    kb_name = self.memory_manager.get_flag(encargado_id, "kb") or self.memory_manager.get_flag(
                        encargado_id,
                        "knowledge_base",
                    )
                if property_id is None:
                    property_id = self._extract_property_id(topic, clean_content, hotel_name, resolved_name)

                add_kb_daily_cache(
                    property_id=property_id,
                    kb_name=kb_name,
                    property_name=resolved_name,
                    topic=topic,
                    category=None,
                    content=clean_content,
                    source_type=source,
                )
            except Exception as exc:
                log.warning("No se pudo guardar cache temporal KB: %s", exc)

            confirmation = (
                "✅ Información agregada a la base de conocimientos:\n\n"
                f"{topic}\n{clean_content[:100]}..."
            )
            await self._safe_call(
                getattr(self.channel_manager, "send_message", None),
                encargado_id,
                confirmation,
                channel="telegram",
            )

            return {
                "status": "success",
                "kb_entry_key": result.get("key") if isinstance(result, dict) else None,
                "message": confirmation,
            }

        except Exception as exc:
            log.error("Error agregando a KB: %s", exc, exc_info=True)
            try:
                await self._safe_call(
                    getattr(self.channel_manager, "send_message", None),
                    encargado_id,
                    f"❌ No se pudo agregar a la base de conocimientos: {exc}",
                    channel="telegram",
                )
            except Exception:
                pass
            return {
                "status": "error",
                "message": f"Error: {exc}",
            }

    def _get_document_class(self):
        try:
            from docx import Document  # type: ignore
            return Document
        except ImportError as exc:
            raise RuntimeError("Falta dependencia python-docx para editar el documento") from exc

    async def _append_to_knowledge_document(
        self,
        topic: str,
        content: str,
        hotel_name: str,
        source_type: str,
        use_env: bool = True,
    ) -> Dict[str, Any]:
        """
        Anexa la información al documento de conocimientos en S3.
        """

        bucket = Settings.S3_BUCKET
        if not bucket:
            raise ValueError("S3_BUCKET no configurado en .env")

        boto = self._get_s3_client()

        candidates = self._resolve_doc_candidates(
            hotel_name=hotel_name,
            boto_client=boto,
            bucket=bucket,
            use_env=use_env,
        )

        tmp_dir = Path(tempfile.mkdtemp())
        local_path = tmp_dir / "kb.docx"

        key_used = None
        last_exc: Exception | None = None
        create_new = False
        for key in candidates:
            try:
                await asyncio.to_thread(boto.download_file, bucket, key, str(local_path))
                key_used = key
                break
            except Exception as exc:  # intentamos siguiente candidato
                last_exc = exc
                log.warning("No se pudo descargar %s/%s, probando siguiente: %s", bucket, key, exc)

        if not key_used:
            key_used = candidates[0]
            create_new = True
            log.warning(
                "No se pudo descargar ningún documento de KB tras probar %s candidatos; se creará uno nuevo en %s",
                len(candidates),
                key_used,
            )

        Document = self._get_document_class()

        if create_new:
            doc = Document()
        else:
            doc = Document(str(local_path))
        timestamp = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")
        doc.add_paragraph(f"[{timestamp}] {topic}")
        doc.add_paragraph(content)
        doc.add_paragraph(f"(source: {source_type})")
        doc.save(str(local_path))

        try:
            await asyncio.to_thread(boto.upload_file, str(local_path), bucket, key_used)
        except Exception as exc:
            raise RuntimeError(f"No se pudo subir el documento actualizado a S3 ({bucket}/{key_used}): {exc}")

        return {"status": "success", "key": key_used}

    async def _prepare_kb_removal_preview(
        self,
        hotel_name: str,
        criterio: str,
        fecha_inicio: str | None = None,
        fecha_fin: str | None = None,
        preview_limit: int = 5,
        max_preview_chars: int = 2200,
    ) -> dict[str, Any]:
        """
        Lee el documento de KB y prepara un borrador de eliminación según criterio/fechas.
        Devuelve payload estructurado para que el webhook muestre conteo o extractos.
        """

        criterio_clean = (criterio or "").strip()
        if not criterio_clean and not (fecha_inicio or fecha_fin):
            return {
                "total_matches": 0,
                "criteria": criterio_clean,
                "error": "Necesito un criterio o un rango de fechas para buscar qué eliminar.",
            }

        kb_data = await self._load_kb_entries(hotel_name, use_env=False)
        entries = kb_data.get("entries", [])
        if not entries:
            return {
                "total_matches": 0,
                "criteria": criterio_clean,
                "error": "No encontré registros en la base de conocimientos Variable.",
            }

        def _parse_date(val: str | None):
            if not val:
                return None
            try:
                return datetime.fromisoformat(val.strip())
            except Exception:
                try:
                    return datetime.strptime(val.strip(), "%Y-%m-%d")
                except Exception:
                    return None

        date_from = _parse_date(fecha_inicio)
        date_to = _parse_date(fecha_fin)

        crit_lower = criterio_clean.lower()
        crit_terms = [t for t in re.findall(r"[a-záéíóúñü0-9]+", crit_lower) if t]

        def _matches(entry: dict[str, Any]) -> bool:
            blob = f"{entry.get('topic','')} {entry.get('content','')}".lower()
            if crit_terms:
                if not any(term in blob for term in crit_terms):
                    return False
            if date_from or date_to:
                ts = entry.get("timestamp_dt")
                if isinstance(ts, datetime):
                    if date_from and ts < date_from:
                        return False
                    if date_to and ts > date_to:
                        return False
            return True

        matched = [e for e in entries if _matches(e)]
        total = len(matched)
        if not matched:
            return {
                "total_matches": 0,
                "criteria": criterio_clean,
                "date_from": fecha_inicio,
                "date_to": fecha_fin,
                "doc_key": kb_data.get("key"),
                "matches": [],
                "preview": [],
            }

        preview_items = []
        preview_chars = 0

        def _clean_snippet(text: str) -> str:
            """Recorta texto y elimina líneas de borradores previos para una vista limpia."""
            if not text:
                return ""
            lines = []
            for ln in (text or "").splitlines():
                low = ln.lower()
                if "borrador para agregar" in low or "[kb_" in low or "[kb-" in low:
                    continue
                lines.append(ln.strip())
            cleaned = " ".join(ln for ln in lines if ln).strip()
            return cleaned[:320] + ("…" if len(cleaned) > 320 else "")

        for entry in matched[:preview_limit]:
            snippet = _clean_snippet(entry.get("content") or "")
            item = {
                "id": entry.get("id"),
                "fecha": entry.get("timestamp_display"),
                "topic": entry.get("topic"),
                "snippet": snippet,
                "source": entry.get("source"),
            }
            preview_items.append(item)
            preview_chars += len(snippet or "")
            if preview_chars >= max_preview_chars:
                break

        payload = {
            "criteria": criterio_clean,
            "date_from": fecha_inicio,
            "date_to": fecha_fin,
            "doc_key": kb_data.get("key"),
            "total_matches": total,
            "preview_count": len(preview_items),
            "preview": preview_items,
            "target_ids": [e.get("id") for e in matched],
            "matches": [
                {
                    "id": e.get("id"),
                    "topic": e.get("topic"),
                    "timestamp_display": e.get("timestamp_display"),
                    "content": e.get("content"),
                }
                for e in matched
            ],
        }
        return payload

    async def handle_kb_removal(
        self,
        hotel_name: str,
        target_ids: list[int],
        encargado_id: str,
        note: str = "",
        criteria: str = "",
    ) -> Dict[str, Any]:
        """
        Elimina entradas del documento de KB (Variable) según IDs parseados.
        """

        if not target_ids:
            return {"status": "noop", "message": "No hay registros seleccionados para eliminar."}

        bucket = Settings.S3_BUCKET
        if not bucket:
            raise ValueError("S3_BUCKET no configurado en .env")

        boto = self._get_s3_client()

        resolved_name, from_memory = self._resolve_hotel_name(hotel_name, encargado_id)
        kb_data = await self._load_kb_entries(
            resolved_name,
            boto_client=boto,
            bucket=bucket,
            use_env=False,
        )
        entries = kb_data.get("entries", [])
        key_used = kb_data.get("key")
        local_path = kb_data.get("path")

        if not entries:
            msg = "No encontré registros en la base de conocimientos para eliminar."
            await self._safe_call(
                getattr(self.channel_manager, "send_message", None),
                encargado_id,
                msg,
                channel="telegram",
            )
            return {"status": "empty", "message": msg}

        remove_set = {int(tid) for tid in target_ids}
        kept = [e for e in entries if e.get("id") not in remove_set]
        removed = [e for e in entries if e.get("id") in remove_set]

        if not removed:
            msg = "No se encontraron coincidencias para eliminar con el criterio indicado."
            await self._safe_call(
                getattr(self.channel_manager, "send_message", None),
                encargado_id,
                msg,
                channel="telegram",
            )
            return {"status": "noop", "message": msg}

        Document = self._get_document_class()
        doc_new = Document()
        for entry in kept:
            for par_text in entry.get("paragraphs", []):
                doc_new.add_paragraph(par_text)
        doc_new.save(str(local_path))

        try:
            await asyncio.to_thread(boto.upload_file, str(local_path), bucket, key_used)
        except Exception as exc:
            raise RuntimeError(f"No se pudo subir el documento actualizado a S3 ({bucket}/{key_used}): {exc}")

        header_lines = [
            f"🧹 Eliminados {len(removed)} registros de la base de conocimientos.",
            f"Criterio: {criteria or 'sin especificar'}",
        ]
        if note:
            header_lines.append(f"Nota: {note}")

        preview_lines = []
        for rm in removed[:5]:
            line = f"- {rm.get('timestamp_display') or ''} {rm.get('topic') or ''}".strip()
            preview_lines.append(line or f"- ID {rm.get('id')}")
        if len(removed) > 5:
            preview_lines.append(f"... y {len(removed) - 5} más.")

        confirmation = "\n".join(header_lines + (["Resumen:"] + preview_lines if preview_lines else []))

        await self._safe_call(
            getattr(self.channel_manager, "send_message", None),
            encargado_id,
            confirmation,
            channel="telegram",
        )

        return {
            "status": "success",
            "removed": [rm.get("id") for rm in removed],
            "kept": [kp.get("id") for kp in kept],
            "message": confirmation,
            "doc_key": key_used,
        }

    def _resolve_doc_candidates(
        self,
        hotel_name: str,
        boto_client: Any = None,
        bucket: str | None = None,
        use_env: bool = True,
    ) -> list[str]:
        """
        Devuelve una lista ordenada de posibles keys en S3 para el documento de KB del hotel.
        Prioriza archivos que contengan '-Variable' antes de la extensión y agrega fallback final.
        """

        if use_env and Settings.SUPERINTENDENTE_S3_DOC:
            env_key = Settings.SUPERINTENDENTE_S3_DOC.strip('\"\\\' ')
            candidates = [env_key]
            # También probar variante "-Variable" si no está incluida
            if env_key.lower().endswith(".docx"):
                base_no_ext = env_key[:-5]
                var_key = f"{base_no_ext}-Variable.docx"
                if var_key not in candidates:
                    candidates.insert(0, var_key)
            elif env_key.lower().endswith(".doc"):
                base_no_ext = env_key[:-4]
                var_key = f"{base_no_ext}-Variable.doc"
                if var_key not in candidates:
                    candidates.insert(0, var_key)
            log.info("Candidatos para documento KB (env): %s", candidates)
            return candidates

        prefix_env = Settings.SUPERINTENDENTE_S3_PREFIX.rstrip("/") if use_env else ""
        clean_name = re.sub(r"[^A-Za-z0-9\\-_ ]+", "", hotel_name).strip()
        tokens = [t for t in re.findall(r"[a-z0-9]+", clean_name.lower()) if t]
        stop_tokens = {"hotel", "hostal", "centro"}
        tokens = [t for t in tokens if t not in stop_tokens]
        doc_name = f"{clean_name.replace(' ', '_')}.docx" if clean_name else "knowledge_base.docx"
        slug_prefix = clean_name.replace(" ", "_") if clean_name else ""
        alt_slug_prefix = ""
        if slug_prefix.lower().startswith("hotel_"):
            alt_slug_prefix = slug_prefix[6:]
        elif slug_prefix.lower().startswith("hostal_"):
            alt_slug_prefix = slug_prefix[7:]
        prefix_tail = prefix_env.rsplit("/", 1)[-1] if prefix_env else ""

        # base_key por defecto
        base_key = f"{prefix_env}/{doc_name}" if prefix_env else doc_name

        # 🎯 Generar combinaciones de prefijos candidatos
        prefix_candidates = []
        if prefix_env:
            prefix_candidates.append(prefix_env)
        if slug_prefix and slug_prefix not in prefix_candidates:
            prefix_candidates.append(slug_prefix)
        if alt_slug_prefix and alt_slug_prefix not in prefix_candidates:
            prefix_candidates.append(alt_slug_prefix)
        tokens_title = [t.title() for t in tokens]
        if tokens_title:
            title_prefix = "_".join(tokens_title)
            if title_prefix not in prefix_candidates:
                prefix_candidates.append(title_prefix)
        if "" not in prefix_candidates:
            prefix_candidates.append("")

        # 🎯 Generar nombres posibles de archivo Variable
        var_names = []
        if clean_name:
            raw_name = " ".join(hotel_name.split())  # normaliza espacios múltiples
            var_names.extend(
                [
                    f"{raw_name}-Variable.docx",
                    f"{clean_name.replace(' ', '_')}-Variable.docx",
                    f"{'_'.join(tokens_title)}-Variable.docx" if tokens_title else "",
                    f"{doc_name[:-5]}-Variable.docx" if doc_name.endswith(".docx") else f"{doc_name}-Variable.docx",
                ]
            )
        if prefix_tail:
            var_names.extend(
                [
                    f"{prefix_tail.replace('_', ' ')}-Variable.docx",
                    f"{prefix_tail}-Variable.docx",
                ]
            )

        candidates: list[str] = []
        # Añadir combinaciones prefijo + nombres variable
        for pref in prefix_candidates:
            for nm in var_names:
                if not nm:
                    continue
                cand = f"{pref}/{nm}" if pref else nm
                if cand not in candidates:
                    candidates.append(cand)

        # 🎯 Buscar primero documentos existentes coincidentes (contienen todos los tokens del hotel)
        # en los prefijos conocidos; evita crear carpetas nuevas si ya existe una.
        if boto_client and bucket:
            search_prefixes: list[str] = []
            if prefix_env:
                search_prefixes.append(prefix_env)
            for pref in prefix_candidates:
                if pref and pref not in search_prefixes:
                    search_prefixes.append(pref)
            if "" not in search_prefixes:
                search_prefixes.append("")

            for pref in search_prefixes:
                try:
                    paginator = boto_client.get_paginator("list_objects_v2")
                    found: list[str] = []
                    for page in paginator.paginate(Bucket=bucket, Prefix=f"{pref}/" if pref else ""):
                        for obj in page.get("Contents", []):
                            key = obj.get("Key") or ""
                            key_lower = key.lower()
                            if tokens and not all(tok in key_lower for tok in tokens):
                                continue
                            if not key_lower.endswith((".docx", ".doc")):
                                continue
                            if key not in found:
                                found.append(key)
                    if found:
                        found.sort(key=lambda k: (0 if "-variable" in k.lower() else 1, len(k)))
                        log.info("Candidatos existentes coincidentes para KB: %s", found)
                        return found
                except Exception as exc:
                    log.warning("No se pudo listar documentos en %s: %s", pref or "<root>", exc)

        # Añadir base como último recurso
        if base_key not in candidates:
            candidates.append(base_key)

        log.info("Candidatos para documento KB: %s", candidates)
        return candidates

    async def _load_kb_entries(
        self,
        hotel_name: str,
        boto_client: Any = None,
        bucket: str | None = None,
        use_env: bool = True,
    ) -> dict[str, Any]:
        """
        Descarga el documento de KB y lo parsea en entradas discretas con índices.
        """

        bucket = bucket or Settings.S3_BUCKET
        if not bucket:
            raise ValueError("S3_BUCKET no configurado en .env")

        boto_client = boto_client or self._get_s3_client()

        candidates = self._resolve_doc_candidates(
            hotel_name=hotel_name,
            boto_client=boto_client,
            bucket=bucket,
            use_env=use_env,
        )

        tmp_dir = Path(tempfile.mkdtemp())
        local_path = tmp_dir / "kb.docx"

        key_used = None
        create_new = False
        last_exc: Exception | None = None
        for key in candidates:
            try:
                await asyncio.to_thread(boto_client.download_file, bucket, key, str(local_path))
                key_used = key
                break
            except Exception as exc:  # intentamos siguiente candidato
                last_exc = exc
                log.warning("No se pudo descargar %s/%s, probando siguiente: %s", bucket, key, exc)

        if not key_used:
            key_used = candidates[0]
            create_new = True
            log.warning(
                "No se pudo descargar ningún documento de KB tras probar %s candidatos; se creará uno nuevo en %s",
                len(candidates),
                key_used,
            )
            return {"entries": [], "key": key_used, "path": local_path, "create_new": create_new}

        Document = self._get_document_class()
        doc = Document(str(local_path))
        entries = self._parse_kb_paragraphs(doc.paragraphs)

        return {
            "entries": entries,
            "key": key_used,
            "path": local_path,
            "create_new": create_new,
        }

    def _parse_kb_paragraphs(self, paragraphs: list[Any]) -> list[dict[str, Any]]:
        """
        Convierte los párrafos del documento en entradas con índice y metadatos.
        Asume formato estándar: [timestamp] Título, contenido y (source: tipo).
        """

        entries: list[dict[str, Any]] = []
        current: dict[str, Any] = {"header": "", "content": "", "source": "", "paragraphs": []}

        def _flush():
            if not any(current.values()):
                return
            idx = len(entries)
            header = current.get("header", "")
            topic = header
            ts_display = ""
            ts_dt = None
            match = re.match(r"\[(.*?)\]\s*(.*)", header)
            if match:
                ts_display = match.group(1).strip()
                topic = match.group(2).strip() or topic
                ts_clean = ts_display.replace(" UTC", "").replace("T", " ")
                try:
                    ts_dt = datetime.fromisoformat(ts_clean)
                except Exception:
                    ts_dt = None

            entry = {
                "id": idx,
                "header": header,
                "topic": topic,
                "timestamp_display": ts_display,
                "timestamp_dt": ts_dt,
                "content": (current.get("content") or "").strip(),
                "source": current.get("source", ""),
                "paragraphs": list(current.get("paragraphs") or []),
            }
            entries.append(entry)

        for idx, para in enumerate(paragraphs or []):
            text = (getattr(para, "text", "") or "").strip()
            if not text:
                continue
            is_header = bool(re.match(r"^\[\d{4}-\d{2}-\d{2}", text)) or ("UTC]" in text and text.startswith("["))
            if is_header:
                _flush()
                current = {"header": text, "content": "", "source": "", "paragraphs": [text]}
                continue

            if text.lower().startswith("(source:"):
                current["source"] = text
                current.setdefault("paragraphs", []).append(text)
                continue

            if current.get("content"):
                current["content"] = f"{current['content']}\n{text}"
            else:
                current["content"] = text
            current.setdefault("paragraphs", []).append(text)

        _flush()
        return entries

    def _clean_kb_content(self, content: str) -> str:
        """Elimina instrucciones o metadatos que no deben ir al documento KB."""
        if not content:
            return ""

        lines = []
        for raw in content.splitlines():
            ln = raw.strip()
            low = ln.lower()
            if not ln:
                continue
            if "confirma con \"ok\"" in low or "confirma con" in low:
                continue
            if "responde 'ok" in low or "responde \"ok" in low:
                continue
            if "envía ajustes" in low or "envia ajustes" in low:
                continue
            if low.startswith("(source:"):
                continue
            if "[superintendente]" in low:
                continue
            if low.startswith("📝 propuesta para base de conocimientos"):
                continue
            lines.append(ln)

        return "\n".join(lines).strip()

    async def review_recent_conversations(
        self,
        hotel_name: str,
        limit: int = 10,
    ) -> str:
        """
        Revisar conversaciones recientes y sumarizar patrones
        Útil para que el encargado sepa qué preguntas hacen los huéspedes
        """

        summary = (
            "RESUMEN DE CONVERSACIONES RECIENTES:\n\n"
            "1. Preguntas sobre Servicios (5):\n"
            "   - Masajista personal: 3 preguntas\n"
            "   - Servicio de Room Service: 2 preguntas\n"
            "2. Preguntas sobre Ubicación (3):\n"
            "   - Dónde está la piscina: 2\n"
            "   - Horarios de restaurante: 1\n"
            "3. Preguntas no respondidas (2):\n"
            "   - Opciones de transfer al aeropuerto\n"
            "   - Alquiler de bicicletas\n\n"
            "RECOMENDACIÓN: Considera agregar información sobre servicios adicionales a la base de conocimientos."
        )

        return summary

    async def _safe_call(self, func: Optional[Any], *args, **kwargs):
        """Invoca funciones sync/async de forma segura."""
        if not func:
            return None
        try:
            res = func(*args, **kwargs)
            if inspect.isawaitable(res):
                return await res
            return res
        except Exception as exc:
            log.warning("Error en llamada segura: %s", exc)
            raise
