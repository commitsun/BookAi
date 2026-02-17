import os
from io import BytesIO
from datetime import datetime, timezone
from typing import Dict, List, Optional, Set
from uuid import uuid4

import boto3
import tiktoken
from dotenv import load_dotenv
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from PyPDF2 import PdfReader
from docx import Document
from supabase import create_client
from pipeline.deadline_filter import filter_expired_sections, is_variable_doc

# =====================================
# 🔧 Cargar configuración
# =====================================
load_dotenv()

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
S3_BUCKET = os.getenv("S3_BUCKET", "bookai-pre-roomdoo")
AWS_REGION = os.getenv("AWS_DEFAULT_REGION", "eu-west-1")

# Clientes globales
supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
s3 = boto3.client("s3", region_name=AWS_REGION)
embeddings_model = OpenAIEmbeddings(model="text-embedding-3-small")
ENCODER = tiktoken.get_encoding("cl100k_base")


# =====================================
# 📝 Utilidades para limpiar y subir documentos
# =====================================
def _docx_bytes_from_text(text: str) -> bytes:
    """Crea un DOCX simple a partir de texto plano separado por párrafos en blanco."""
    doc = Document()
    blocks = [b.strip() for b in text.split("\n\n")] if text else []
    for block in blocks if blocks else [""]:
        doc.add_paragraph(block)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _upload_cleaned_doc(key: str, text: str) -> None:
    """Reescribe en S3 el documento Variable ya depurado."""
    ext = os.path.splitext(key)[-1].lower()
    try:
        if ext == ".docx":
            data = _docx_bytes_from_text(text)
            s3.put_object(Bucket=S3_BUCKET, Key=key, Body=data)
        elif ext == ".txt":
            s3.put_object(Bucket=S3_BUCKET, Key=key, Body=text.encode("utf-8"))
        else:
            print(f"⚠️ Tipo no soportado para reescritura: {ext}")
            return
        print(f"📤 Documento limpiado subido a S3: {key}")
    except Exception as exc:
        print(f"⚠️ No se pudo reescribir {key} en S3: {exc}")


# =====================================
# 📄 Lectura y chunking de documentos
# =====================================
def load_text_from_s3(key: str) -> str:
    """
    Descarga y lee el contenido de un archivo desde S3.
    Soporta: .txt, .docx, .pdf
    """
    obj = s3.get_object(Bucket=S3_BUCKET, Key=key)
    raw_data = obj["Body"].read()
    ext = os.path.splitext(key)[-1].lower()

    if ext == ".txt":
        return raw_data.decode("utf-8", errors="ignore")
    if ext == ".docx":
        doc = Document(BytesIO(raw_data))
        return "\n".join([p.text for p in doc.paragraphs])
    if ext == ".pdf":
        pdf = PdfReader(BytesIO(raw_data))
        return "\n".join([page.extract_text() or "" for page in pdf.pages])

    raise ValueError(f"❌ Tipo de archivo no soportado: {ext}")


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Divide el texto en chunks con tamaño por tokens para mejorar recuperación."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        separators=["\n\n", "\n", ".", " "],
        length_function=lambda s: len(ENCODER.encode(s)),
    )
    return splitter.split_text(text)


def list_s3_files(prefix: str) -> List[Dict[str, str]]:
    """Lista archivos dentro de un prefijo y devuelve metadatos básicos para detectar cambios."""
    files: List[Dict[str, str]] = []
    paginator = s3.get_paginator("list_objects_v2")

    normalized_prefix = prefix if prefix.endswith("/") else f"{prefix}/"

    for page in paginator.paginate(Bucket=S3_BUCKET, Prefix=normalized_prefix):
        for obj in page.get("Contents", []):
            key = obj["Key"]
            if key.endswith("/"):
                continue
            files.append(
                {
                    "key": key,
                    "file_name": os.path.basename(key),
                    "etag": obj.get("ETag", "").strip('"'),
                    "last_modified": obj.get("LastModified").isoformat()
                    if obj.get("LastModified")
                    else "",
                }
            )

    return files


# =====================================
# 🧠 Gestión de Supabase (estado de archivos)
# =====================================
def fetch_existing_file_etag(
    table_name: str, source_key: str, legacy_source: Optional[str] = None
) -> Optional[str]:
    """Devuelve el etag almacenado para una fuente ya vectorizada (None si no existe)."""
    try:
        response = (
            supabase.table(table_name)
            .select("metadata")
            .eq("metadata->>source_key", source_key)
            .limit(1)
            .execute()
        )
    except Exception as exc:
        print(f"⚠️ No se pudo leer estado previo de {source_key}: {exc}")
        return None

    if not response.data and legacy_source:
        try:
            response = (
                supabase.table(table_name)
                .select("metadata")
                .eq("metadata->>source", legacy_source)
                .limit(1)
                .execute()
            )
        except Exception as exc:
            print(f"⚠️ No se pudo leer estado previo legacy de {legacy_source}: {exc}")
            return None

    if not response.data:
        return None

    metadata = response.data[0].get("metadata") or {}
    return metadata.get("etag")


def list_vectorized_sources(table_name: str) -> Set[str]:
    """Obtiene el listado de archivos ya vectorizados (solo nombre)."""
    try:
        response = supabase.table(table_name).select("metadata").execute()
    except Exception as exc:
        print(f"⚠️ No se pudieron listar embeddings previos: {exc}")
        return set()

    sources: Set[str] = set()
    for row in response.data or []:
        meta = row.get("metadata") or {}
        source_key = meta.get("source_key")
        source = meta.get("source")
        if source_key:
            sources.add(source_key)
        elif source:
            sources.add(source)
    return sources


def delete_file_from_supabase(
    table_name: str, source_key: str, legacy_source: Optional[str] = None
) -> None:
    """Elimina los embeddings de un archivo concreto sin tocar el resto de la tabla."""
    deleted_total = 0
    try:
        response = (
            supabase.table(table_name)
            .delete()
            .eq("metadata->>source_key", source_key)
            .execute()
        )
        deleted = len(response.data or []) if hasattr(response, "data") else 0
        deleted_total += deleted
    except Exception as exc:
        print(f"⚠️ No se pudo eliminar {source_key} de {table_name}: {exc}")

    if legacy_source:
        try:
            response = (
                supabase.table(table_name)
                .delete()
                .eq("metadata->>source", legacy_source)
                .execute()
            )
            deleted = len(response.data or []) if hasattr(response, "data") else 0
            deleted_total += deleted
        except Exception as exc:
            print(f"⚠️ No se pudo eliminar legacy {legacy_source} de {table_name}: {exc}")

    print(f"🗑️  Eliminados {deleted_total} registros de {source_key} en {table_name}.")


def purge_table(table_name: str) -> None:
    """Borra todos los embeddings de la tabla sin eliminar la tabla."""
    try:
        supabase.table(table_name).delete().neq("id", None).execute()
        print(f"🧹 Tabla {table_name} limpiada por completo.")
    except Exception as exc:
        print(f"⚠️ No se pudo limpiar {table_name}: {exc}")


# =====================================
# 🧠 Vectorización e inserción en Supabase
# =====================================
def save_chunks_to_supabase(
    table_name: str,
    doc_name: str,
    source_key: str,
    chunks: List[str],
    *,
    etag: Optional[str],
    last_modified: Optional[str],
) -> None:
    """Vectoriza e inserta los chunks en la tabla del hotel, conservando el orden y guardando etag."""
    print(f"🧩 Generando embeddings para {doc_name}...")

    vectors = embeddings_model.embed_documents(chunks)

    rows = [
        {
            "id": str(uuid4()),
            "position": idx,
            "content": chunk,
            "embedding": vector,
            "metadata": {
                "source": doc_name,
                "source_key": source_key,
                "chunk": idx,
                "etag": etag,
                "last_modified": last_modified,
                "indexed_at": datetime.now(timezone.utc).isoformat(),
            },
        }
        for idx, (chunk, vector) in enumerate(zip(chunks, vectors))
    ]

    supabase.table(table_name).insert(rows).execute()
    print(f"✅ Insertados {len(rows)} chunks en {table_name} (orden preservado)")


# =====================================
# 🚀 Vectorización por hotel (incremental)
# =====================================
def vectorize_hotel_docs(hotel_folder: str, *, full_refresh: bool = False) -> None:
    """
    Vectoriza todo lo que haya en la carpeta del hotel:
    - Revectoriza cada archivo (borra lo previo de ese archivo antes de insertar)
    - Elimina de Supabase los archivos que ya no están en S3
    """
    table_name = f"kb_{os.path.basename(hotel_folder).lower()}"
    print(f"\n🚀 Iniciando vectorización para: {table_name}")

    prefix = f"{hotel_folder}/"
    if full_refresh:
        purge_table(table_name)

    s3_files = list_s3_files(prefix)
    if not s3_files:
        print(f"⚠️ No se encontraron archivos en {prefix}")
        return

    print(f"📂 Archivos detectados en S3 ({len(s3_files)}): {[f['file_name'] for f in s3_files]}")

    current_sources = {f["key"] for f in s3_files}
    vectorized_sources = list_vectorized_sources(table_name)

    # 1) Limpiar embeddings de archivos que ya no existen en S3
    missing_sources = vectorized_sources - current_sources
    for source in sorted(missing_sources):
        print(f"➖ {source} ya no existe en S3, se elimina de Supabase.")
        delete_file_from_supabase(table_name, source)

    # 2) Procesar nuevos o modificados
    for file_info in s3_files:
        source_key = file_info["key"]
        file_name = file_info["file_name"]
        etag = file_info["etag"]
        last_modified = file_info["last_modified"]

        existing_etag = fetch_existing_file_etag(
            table_name,
            source_key,
            legacy_source=file_name,
        )
        if existing_etag and existing_etag == etag:
            print(f"⏭️  Sin cambios en {source_key}, se omite revectorización.")
            continue

        # Solo revectorizamos cuando hay cambios (etag distinto) o primera carga.
        delete_file_from_supabase(table_name, source_key, legacy_source=file_name)
        print(f"🔄 Revectorizando {file_name} ...")

        try:
            text = load_text_from_s3(file_info["key"])

            if is_variable_doc(file_name):
                filtered_text, removed = filter_expired_sections(text)
                if removed:
                    print(f"⏳ {file_name}: se eliminaron {len(removed)} bloques caducados:")
                    for detail in removed:
                        print(f"   - {detail}")
                    _upload_cleaned_doc(file_info["key"], filtered_text)
                text = filtered_text
                if not text.strip():
                    print(f"⚠️ {file_name} quedó vacío tras limpiar fechas vencidas, se omite.")
                    continue

            chunks = chunk_text(text)
            if not chunks:
                print(f"⚠️ {file_name} no tiene contenido legible tras limpieza, se omite.")
                continue
            save_chunks_to_supabase(
                table_name,
                file_name,
                source_key,
                chunks,
                etag=etag,
                last_modified=last_modified,
            )
        except Exception as exc:
            print(f"⚠️ Error procesando {file_name}: {exc}")

    print(f"🎉 Vectorización completada para {table_name} ✅")


# =====================================
# ▶️ CLI
# =====================================
if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("❌ Uso: python -m pipeline.vectorizer <nombre_carpeta_hotel> [--full-refresh]")
    else:
        folder = sys.argv[1]
        full_refresh_flag = "--full-refresh" in sys.argv[2:]
        vectorize_hotel_docs(folder, full_refresh=full_refresh_flag)
